# -*- coding: utf-8 -*-
"""
ChatGPT conversation -> Word converter
======================================
Supports three sources:
  1. Web-page HTML saved with Ctrl+S
  2. The officially exported conversations.json (full export)
  3. Bookmark-scraped {title, messages:[{role,text}]} JSON

Usage (put the downloaded json on drive D and tell me, or run it yourself):
    python chatgpt_to_word.py "D:\\chatgpt_dom_export.json"
    python chatgpt_to_word.py "D:\\chatgpt_dom_export (1).json" "D:\\另一个.html"
"""
import argparse
import glob
import html
import json
import os
import re
import sys
from html.parser import HTMLParser

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

VOID_TAGS = {
    "br", "hr", "img", "input", "meta", "link", "source", "wbr",
    "area", "base", "col", "embed", "track", "path", "circle",
    "polygon", "polyline", "rect", "use", "ellipse", "line",
}


class ChatGPTParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.messages = []
        self._cur_role = None
        self._cur_parts = []
        self._depth = 0

    def _finish_current(self):
        if self._cur_role is not None:
            self.messages.append((self._cur_role, "".join(self._cur_parts)))
            self._cur_role = None
            self._cur_parts = []
            self._depth = 0

    def handle_starttag(self, tag, attrs):
        d = dict(attrs)
        if d.get("data-message-author-role") in ("user", "assistant"):
            self._finish_current()
            self._cur_role = d["data-message-author-role"]
            self._cur_parts = []
            self._depth = 1
            return
        if self._cur_role is not None:
            if tag not in VOID_TAGS:
                self._depth += 1
            if tag in ("br", "p", "li", "pre", "blockquote"):
                self._cur_parts.append("\n")
            if tag == "li":
                self._cur_parts.append("• ")

    def handle_endtag(self, tag):
        if self._cur_role is not None:
            if tag in VOID_TAGS:
                return
            self._depth -= 1
            if tag in ("p", "pre", "blockquote", "h1", "h2", "h3", "li"):
                self._cur_parts.append("\n")
            if self._depth <= 0:
                self._finish_current()

    def handle_data(self, data):
        if self._cur_role is not None:
            self._cur_parts.append(data)


# Junk lines to strip from exported HTML. The Chinese alternatives
# ("ChatGPT 说", "更多", "深度研究", "推理", "编写代码", "语音", ...) are runtime
# matching data for the Chinese ChatGPT UI and MUST stay verbatim.
JUNK_LINES = re.compile(
    r"^\s*(Copy code|Copy|Edit|ChatGPT|ChatGPT 说|4o|o1|o3|GPT-\S*|Stop generating|"
    r"Regenerate|更多|Listen|Read aloud|Search|深度研究|推理|编写代码|语音)\s*$",
    re.IGNORECASE)


def clean_text(text: str) -> str:
    text = html.unescape(text)
    lines = [ln.rstrip() for ln in text.splitlines()]
    lines = [ln for ln in lines if not JUNK_LINES.match(ln)]
    out, prev_blank = [], False
    for ln in lines:
        blank = not ln.strip()
        if blank and prev_blank:
            continue
        out.append(ln)
        prev_blank = blank
    return "\n".join(out).strip()


def parse_chatgpt_html(path: str):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    content = re.sub(r"<script[\s\S]*?</script>", "", content, flags=re.I)
    content = re.sub(r"<style[\s\S]*?</style>", "", content, flags=re.I)
    parser = ChatGPTParser()
    parser.feed(content)
    messages = []
    for role, text in parser.messages:
        text = clean_text(text)
        if text:
            messages.append((role, text))
    return messages


def _node_text(message: dict) -> str:
    parts = []
    for p in message.get("content", {}).get("parts", []):
        if isinstance(p, str):
            parts.append(p)
        elif isinstance(p, dict):
            if p.get("content_type") == "text":
                parts.append(p.get("text", ""))
            else:
                parts.append(f"[{p.get('content_type', 'attachment')}]")
    return clean_text("\n".join(parts))


def parse_conversation(conv: dict):
    mapping = conv.get("mapping", {})
    nodes = []
    for node in mapping.values():
        msg = node.get("message")
        if not msg or msg.get("role") not in ("user", "assistant"):
            continue
        nodes.append((msg.get("create_time") or 0, msg))
    nodes.sort(key=lambda x: x[0])
    out = []
    for _, msg in nodes:
        text = _node_text(msg)
        if text:
            out.append((msg["role"], text))
    return out


def parse_conversations_json(path: str):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        data = json.load(f)
    result = []
    for conv in data:
        title = (conv.get("title") or "Untitled conversation").strip()
        msgs = parse_conversation(conv)
        if msgs:
            result.append((title, msgs))
    result.sort(key=lambda x: x[0])
    return result


def parse_dom_json(path: str):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        data = json.load(f)
    if isinstance(data, list):
        return parse_conversations_json(path)
    title = (data.get("title") or "ChatGPT conversation").strip()
    msgs = []
    for m in data.get("messages", []):
        role = m.get("role", "assistant")
        text = clean_text(m.get("text", ""))
        if text and role in ("user", "assistant"):
            msgs.append((role, text))
    return [(title, msgs)] if msgs else []


# Display labels for the two roles in the generated document.
# "Me" renders the user's side; "ChatGPT" renders the assistant's side.
ROLE_NAME = {"user": "Me", "assistant": "ChatGPT"}
ROLE_COLOR = {"user": RGBColor(0x1F, 0x4E, 0x79),
              "assistant": RGBColor(0x0B, 0x6E, 0x4F)}


def build_docx(messages, title: str):
    doc = Document()
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Message count: {len(messages)}    Generated by: chatgpt_to_word")
    doc.add_paragraph("─" * 60)
    for i, (role, text) in enumerate(messages, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {ROLE_NAME.get(role, role)}:")
        run.bold = True
        run.font.color.rgb = ROLE_COLOR.get(role, RGBColor(0, 0, 0))
        run.font.size = Pt(12)
        for ln in text.split("\n"):
            para = doc.add_paragraph()
            para.add_run(ln if ln else " ")
        doc.add_paragraph()
    return doc


def unique_path(out_dir: str, name: str) -> str:
    base, ext = os.path.splitext(name)
    p = os.path.join(out_dir, name)
    i = 1
    while os.path.exists(p):
        p = os.path.join(out_dir, f"{base}_{i}{ext}")
        i += 1
    return p


def safe_filename(title: str) -> str:
    title = re.sub(r'[\\/:*?"<>|\r\n\t]', "_", title).strip()
    title = title.strip(" ._")
    return title[:80] or "Untitled conversation"


def convert_one(path: str, out_dir: str, counter: dict):
    ext = os.path.splitext(path)[1].lower()
    if ext in (".html", ".htm"):
        msgs = parse_chatgpt_html(path)
        if not msgs:
            print(f"  [skip] No conversation parsed: {path}")
            return
        base = os.path.splitext(os.path.basename(path))[0]
        out_path = unique_path(out_dir, f"{base}.docx")
        build_docx(msgs, base or "ChatGPT conversation").save(out_path)
        counter["n"] += 1
        counter["msgs"] += len(msgs)
        print(f"  [OK] -> {out_path} ({len(msgs)} messages)")
        return
    if ext == ".json":
        convs = parse_dom_json(path)
        if not convs:
            print(f"  [skip] No conversation parsed: {path}")
            return
        for title, msgs in convs:
            fname = safe_filename(title)
            out_path = unique_path(out_dir, f"{fname}.docx")
            build_docx(msgs, title).save(out_path)
            counter["n"] += 1
            counter["msgs"] += len(msgs)
            print(f"  [OK] \"{title}\" ({len(msgs)} messages) -> {out_path}")
        return
    print(f"  [skip] Unsupported type: {path}")


def main():
    ap = argparse.ArgumentParser(description="ChatGPT conversation -> Word")
    ap.add_argument("files", nargs="*", help="HTML / JSON files")
    ap.add_argument("--dir", "-d", help="process all .html/.json in a directory")
    ap.add_argument("--out", "-o", default=r"D:\ChatGPT对话导出", help="output directory")
    args = ap.parse_args()

    files = list(args.files)
    if args.dir:
        for ext in ("*.html", "*.htm", "*.json"):
            files.extend(sorted(glob.glob(os.path.join(args.dir, ext))))
    if not files:
        ap.print_help()
        sys.exit(1)

    os.makedirs(args.out, exist_ok=True)
    counter = {"n": 0, "msgs": 0}
    print(f"Output directory: {args.out}")
    for f in files:
        convert_one(f, args.out, counter)
    print(f"\nDone: {counter['n']} Word document(s), {counter['msgs']} messages in total.")


if __name__ == "__main__":
    main()
